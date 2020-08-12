import requests
from bs4 import BeautifulSoup
import sys
import datetime
import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
import cnnScraper, jazeeraScraper, forbesScraper, foxScraper, washingtonScraper, nytScraper, cnbcScraper, hillScraper, nprScraper, politicoScraper, guardianScraper, reutersScraper, usatodayScraper

URL = 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRFZxYUdjU0FtVnVHZ0pWVXlnQVAB?hl=en-US&gl=US&ceid=US%3Aen'
page = requests.get(URL)
soup = BeautifulSoup(page.content, 'html.parser')

content = soup.find('main', class_= 'HKt8rc CGNRMc')
articles = content.find_all('div', class_= 'NiLAwe y6IFtc R7GTQ keNKEd j7vNaf nID9nc')
i = 1
numberArticles = int(input("How many articles do you want?"))
NEWS_SOURCES = ['CNN', 'CNBC', 'Al Jazeera', 'Al Jazeera English', 'Reuters', 'Fox News', 'The Guardian', 'NPR', 'The Washington Post', 'The New York Times', 'The Hill', 'POLITICO', 'Forbes', 'USA TODAY']
NORMAL_SOURCES = ['CNN', 'Al Jazeera', 'Al Jazeera English', 'Reuters', 'Fox News', 'The Guardian', 'NPR', 'The New York Times', 'The Hill', 'POLITICO', 'Forbes', 'USA TODAY']

date = str(datetime.datetime.now())
docdate = date.split(' ')[0]
document = Document()
h = document.add_heading(level=0)
documentTitle = h.add_run(f'News Articles {docdate}')
h.alignment = 1
documentTitle.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

for article in articles:
    findSource = article.find('a', class_='wEwyrc AVN2gc uQIVzc Sksgp')
    source = findSource.text
    findTitle = article.find('a', class_='DY5T1d')
    title = findTitle.text
    findLink = article.find('a', href=True)
    if source.strip() in NEWS_SOURCES:
        print(title + ' | ' + source)
        choice = input("Do you want this article [y/n]: ")
        if choice.strip() == 'y':
            link = 'https://news.google.com/' + findLink['href'].strip('./')
            if source.strip() == 'CNN':
                paragraphs = cnnScraper.scrapeCNN(link)
            if (source.strip() == 'Al Jazeera English') or (source.strip() == 'Al Jazeera'):
                paragraphs = jazeeraScraper.scrapeJazeera(link)
            if source.strip() == 'Forbes':
                paragraphs = forbesScraper.scrapeForbes(link)
            if source.strip() == 'Fox News':
                paragraphs = foxScraper.scrapeFox(link)
            if source.strip() == 'The Washington Post':
                articleType = len(washingtonScraper.scrapeWashington(link))
            if source.strip() == 'The New York Times':
                paragraphs = nytScraper.scrapeNYT(link)
            if source.strip() == 'CNBC':
                paragraphs = cnbcScraper.scrapeCNBC(link)
            if source.strip() == 'The Hill':
                paragraphs = hillScraper.scrapeHill(link)
            if source.strip() == 'NPR':
                paragraphs = nprScraper.scrapeNPR(link)
            if source.strip() == 'POLITICO':
                paragraphs = politicoScraper.scrapePolitico(link)
            if source.strip() == 'The Guardian':
                paragraphs = guardianScraper.scrapeGuardian(link)
            if source.strip() == 'Reuters':
                paragraphs = reutersScraper.scrapeReuters(link)
            if source.strip() == 'USA TODAY':
                paragraphs = usatodayScraper.scrapeUSA(link)
            
            h = document.add_heading(level=2)
            heading = h.add_run(f'{title} | {source}')
            if i != 1:
                h.paragraph_format.page_break_before = True
            h.paragraph_format.space_after = Pt(18)
            heading.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

            if source.strip() == 'The Washington Post':
                if articleType == 1:
                    paragraphs = washingtonScraper.scrapeWashington(link)
                    for paragraph in paragraphs:
                        p = document.add_paragraph()
                        p.alignement = 1
                        para = p.add_run(paragraph.text)
                        para.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
                elif articleType == 2:
                    teaser, paragraphs = washingtonScraper.scrapeWashington(link)
                    t = document.add_paragraph()
                    t.alignement = 1
                    tea = t.add_run(teaser.text)
                    tea.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
                    for paragraph in paragraphs:
                        p = document.add_paragraph()
                        p.alignement = 1
                        para = p.add_run(paragraph.text)
                        para.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
            
            if source.strip() == 'CNBC':
                try:
                    for group in paragraphs:
                        for paragraph in group.find_all('p'):
                            p = document.add_paragraph()
                            p.alignement = 1
                            para = p.add_run(paragraph.text)
                            para.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
                except:
                    pass
            
            elif source.strip() in NORMAL_SOURCES:
                for paragraph in paragraphs:
                    p = document.add_paragraph()
                    p.alignement = 1
                    para = p.add_run(paragraph.text)
                    para.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    else:
        continue
    if i == numberArticles:
        break
    i += 1

doctitle = f'News_Articles_{docdate}.docx'
document.save(doctitle)
if sys.platform == 'darwin':
    downloadUser = input('Please enter your Mac username (case sensitive): ')
    document.save(f'/Users/{downloadUser}/Desktop/{doctitle}')
if sys.platform == 'win32':
    downloadUser = input('Please enter your windows username (case sensitive): ')
    document.save('C:/Users/%s/Desktop/%s' % (downloadUser, doctitle))
print(f'{doctitle} has been downloaded with {str(numberArticles)} articles on your desktop')